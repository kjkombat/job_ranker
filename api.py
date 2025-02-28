import asyncio
import json
from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.responses import JSONResponse
from fastapi.openapi.docs import get_swagger_ui_html
from fastapi.staticfiles import StaticFiles
import os
import tempfile
import pdfplumber
from typing import List, Optional
import docx
from pydantic import BaseModel, Field
from llm import llm_response, llm_generate_score
import pandas as pd

# import pandas as pd

app = FastAPI(
    title="Criteria Extraction API",
    description="Extract criteria from job descriptions",
    swagger_ui=True,
)


class CriteriaResponse(BaseModel):
    status: str = Field(..., example="success", description="Status of the operation")
    message: str = Field(
        ...,
        example="Successfully processed file: job_description.pdf",
        description="Operation result message",
    )
    criteria: List[str] = Field(
        ...,
        example=[
            "5+ years of Python experience",
            "Strong knowledge of Django",
            "Bachelor's degree in Computer Science",
        ],
        description="Extracted criteria from the job description",
    )


class ScoreResponse(BaseModel):
    status: str = Field(..., example="success", description="Status of the operation")
    csv: str = Field(
        ...,
        example="username,Python experience,Django knowledge,Education,total\nJohn Doe,5,4,3,12\nJane Smith,3,5,4,12",
        description="CSV-formatted string with scoring results",
    )


@app.post(
    "/extract-criteria",
    response_model=CriteriaResponse,
    tags=["Job Descriptions"],
    summary="Extract criteria from a job description",
    include_in_schema=True,
    description="""
    Upload a job description document (PDF or DOCX) to extract key hiring criteria.
    The system will analyze the document and return a list of important requirements.
    """,
    responses={
        200: {
            "description": "Successful criteria extraction",
            "content": {
                "application/json": {
                    "example": {
                        "status": "success",
                        "message": "Successfully processed file: job_description.pdf",
                        "criteria": [
                            "5+ years of Python experience",
                            "Strong knowledge of Django",
                            "Bachelor's degree in Computer Science",
                        ],
                    }
                }
            },
        },
        400: {
            "description": "Bad request - Invalid file type",
            "content": {
                "application/json": {
                    "example": {"detail": "Only PDF and DOCX files are supported"}
                }
            },
        },
        500: {
            "description": "Server error during processing",
            "content": {
                "application/json": {
                    "example": {
                        "detail": "An error occurred during processing: [error details]"
                    }
                }
            },
        },
    },
)
async def extract_criteria(file: UploadFile = File(...)):
    """
    Extract criteria from a job description file (PDF or DOCX)

    Args:
        file: Uploaded job description file (PDF or DOCX)

    Returns:
        JSON response indicating success and message
    """
    # Validate file type
    if not file.filename.endswith((".pdf", ".docx")):
        raise HTTPException(
            status_code=400, detail="Only PDF and DOCX files are supported"
        )

    try:
        with tempfile.NamedTemporaryFile(
            delete=False, suffix=os.path.splitext(file.filename)[1]
        ) as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_file_path = temp_file.name

        # Extract text based on file type
        file_extension = os.path.splitext(file.filename)[1].lower()

        if file_extension == ".pdf":
            # Extract text from PDF using pdfplumber
            with pdfplumber.open(temp_file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    extracted_text = page.extract_text()
                    if extracted_text:  # Handle potential None returns
                        text += extracted_text + "\n"

        elif file_extension == ".docx":
            # Extract text from DOCX using python-docx
            doc = docx.Document(temp_file_path)
            paragraphs_text = [paragraph.text for paragraph in doc.paragraphs]
            text = "\n".join(paragraphs_text)

        criteria = await llm_response(text)

        os.unlink(temp_file_path)

        return JSONResponse(
            status_code=200,
            content={
                "status": "success",
                "message": f"Successfully processed file: {file.filename}",
                "criteria": criteria.criteria,
            },
        )

    except Exception as e:
        if "temp_file_path" in locals():
            os.unlink(temp_file_path)

        raise HTTPException(
            status_code=500, detail=f"An error occurred during processing: {str(e)}"
        )


@app.post(
    "/score-resumes",
    response_model=ScoreResponse,
    tags=["Resume Scoring"],
    summary="Score multiple resumes against criteria",
    include_in_schema=True,
    description="""
    Upload multiple resumes (PDF or DOCX) and score them against specified criteria.
    The system will analyze each resume and assign scores from 1-5 for each criterion.
    
    The response includes a CSV with scores for each resume and criterion, plus a total score.
    """,
    responses={
        200: {
            "description": "Successful resume scoring",
            "content": {
                "application/json": {
                    "example": {
                        "status": "success",
                        "csv": "username,Python experience,Django knowledge,Education,total\nJohn Doe,5,4,3,12\nJane Smith,3,5,4,12",
                    }
                }
            },
        },
        400: {
            "description": "Bad request - Invalid files or criteria",
            "content": {
                "application/json": {
                    "example": {
                        "detail": "Criteria must be a JSON-encoded list of strings"
                    }
                }
            },
        },
        500: {
            "description": "Server error during processing",
            "content": {
                "application/json": {
                    "example": {
                        "detail": "An error occurred during processing: [error details]"
                    }
                }
            },
        },
    },
)
async def score_resumes(files: List[UploadFile] = File(...), criteria: str = Form()):
    """
    Score resumes against specified criteria

    Args:
        criteria: JSON string containing list of criteria
        files: List of uploaded resume files (PDF or DOCX)

    Returns:
        JSON response with scores for each resume
    """
    try:
        # Parse the criteria JSON string into a Python list
        criteria_list = json.loads(criteria)

        if not isinstance(criteria_list, list):
            raise HTTPException(
                status_code=400,
                detail="Criteria must be a JSON-encoded list of strings",
            )

        if len(criteria_list) < 1:
            raise HTTPException(
                status_code=400,
                detail="At least one criteria must be provided",
            )

        # Validate files
        for file in files:
            if not file.filename.endswith((".pdf", ".docx")):
                raise HTTPException(
                    status_code=400,
                    detail=f"File {file.filename} is not supported. Only PDF and DOCX files are accepted.",
                )

        # List to store results for each resume
        list_of_scores = []
        tasks = [process_file(file, criteria_list) for file in files]
        list_of_scores = await asyncio.gather(*tasks)

        # Process each dictionary calculate totals
        for score_dict in list_of_scores:
            total_score = 0

            for key, value in list(score_dict.items()):
                # Skip the username field
                if key != "username":
                    if isinstance(value, str):
                        try:
                            converted_value = int(value)
                            if converted_value >= 1 and converted_value <= 5:
                                total_score += converted_value
                            else:
                                score_dict[key] = "Error"
                        except ValueError:
                            # If conversion fails, keep as string but don't add to total
                            pass
                    elif isinstance(value, (int, float)):
                        if value >= 1 and value <= 5:
                            total_score += value
                        else:
                            score_dict[key] = "Error"

            # Add the total score to the dictionary
            score_dict["total"] = total_score

        # Convert to DataFrame and then to CSV
        output_csv = pd.DataFrame(list_of_scores).to_csv(index=False)

        return JSONResponse(
            status_code=200,
            content={
                "status": "success",
                "csv": output_csv,
            },
        )
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"An error occurred during processing: {str(e)}"
        )


async def process_file(file: UploadFile, criteria_list: List[str]):
    """Process a single file and return its scores."""
    with tempfile.NamedTemporaryFile(
        delete=False, suffix=os.path.splitext(file.filename)[1]
    ) as temp_file:
        content = await file.read()
        temp_file.write(content)
        temp_file_path = temp_file.name

    try:
        file_extension = os.path.splitext(file.filename)[1].lower()

        if file_extension == ".pdf":
            # Extract text from PDF using pdfplumber
            with pdfplumber.open(temp_file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    extracted_text = page.extract_text()
                    if extracted_text:  # Handle potential None returns
                        text += extracted_text + "\n"

        elif file_extension == ".docx":
            # Extract text from DOCX using python-docx
            doc = docx.Document(temp_file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            # Also extract text from tables in the document
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        # Generate score using LLM
        scores = await llm_generate_score(text, criteria_list)
        return dict(scores)
    finally:
        # Ensure temp file is always deleted
        if os.path.exists(temp_file_path):
            os.unlink(temp_file_path)


# Add a root endpoint with documentation
@app.get("/", include_in_schema=False)
async def root():
    return {
        "message": "Welcome to the Resume Scoring API",
        "documentation": "/docs",
        "endpoints": [
            {
                "path": "/extract-criteria",
                "method": "POST",
                "description": "Extract criteria from job descriptions",
            },
            {
                "path": "/score-resumes",
                "method": "POST",
                "description": "Score resumes against criteria",
            },
        ],
    }


@app.get("/docs", include_in_schema=False)
async def custom_swagger_ui_html():
    return get_swagger_ui_html(
        openapi_url=app.openapi_url,
        title=app.title + " - API Documentation",
        swagger_js_url="https://cdn.jsdelivr.net/npm/swagger-ui-dist@4/swagger-ui-bundle.js",
        swagger_css_url="https://cdn.jsdelivr.net/npm/swagger-ui-dist@4/swagger-ui.css",
        swagger_favicon_url="https://fastapi.tiangolo.com/img/favicon.png",
    )
